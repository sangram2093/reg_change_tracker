import os
import json
import re
from flask import Flask, request, render_template, redirect, url_for, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
from vertex_llm import init_vertexai, get_summary_with_context, get_entity_relationship_with_context
from utils import extract_text_from_pdf, parse_graph_data, generate_kop
from db_models import db, Regulation, Upload, Summary, EntityGraph
from docx import Document
from io import BytesIO

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'postgresql+psycopg2://db_admin:pwd@127.0.0.1:5432/postgres'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db.init_app(app)

init_vertexai()

@app.route("/", methods=["GET", "POST"])
def index():
    regulations = Regulation.query.all()
    if request.method == "POST":
        regulation_id = request.form['regulation']
        old_path = request.form['old_path'].strip()
        new_path = request.form['new_path'].strip()

        if not os.path.exists(old_path):
            return "Old PDF path is invalid.", 400
        if new_path and not os.path.exists(new_path):
            return "New PDF path is invalid.", 400

        upload = Upload(regulation_id=regulation_id, old_path=old_path, new_path=new_path)
        db.session.add(upload)
        db.session.commit()

        process_upload(upload.id)
        return redirect(url_for("compare", upload_id=upload.id))

    return render_template("index.html", regulations=regulations)

def markdown_to_docx(doc: Document, text: str):
    lines = text.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("###"):
            doc.add_heading(stripped.lstrip("#").strip(), level=3)
        elif stripped.startswith("##"):
            doc.add_heading(stripped.lstrip("#").strip(), level=2)
        elif stripped.startswith("#"):
            doc.add_heading(stripped.lstrip("#").strip(), level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', stripped):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', stripped), style='List Number')
        elif "**" in stripped:
            para = doc.add_paragraph()
            while "**" in stripped:
                before, bold, rest = stripped.split("**", 2)
                para.add_run(before)
                bold_run = para.add_run(bold)
                bold_run.bold = True
                stripped = rest
            para.add_run(stripped)
        else:
            doc.add_paragraph(stripped)

def process_upload(upload_id):
    upload = Upload.query.get(upload_id)

    # Clean existing entries for this upload ID
    db.session.query(Summary).filter_by(upload_id=upload.id).delete()
    db.session.query(EntityGraph).filter_by(upload_id=upload.id).delete()

    if not upload.old_path:
        # First-time upload
        new_text = extract_text_from_pdf(upload.new_path)
        new_summary = get_summary_with_context(new_text)
        new_json = get_entity_relationship_with_context(new_summary)

        G_new = parse_graph_data(json.loads(new_json))
        graph_new_json = json.dumps({
            "nodes": [{"id": n, **G_new.nodes[n]} for n in G_new.nodes],
            "edges": [{"from": u, "to": v, **G_new[u][v]} for u, v in G_new.edges]
        })

        db.session.add(Summary(upload_id=upload.id, old_summary=None, new_summary=new_summary))
        db.session.add(EntityGraph(
            upload_id=upload.id,
            old_json=None,
            new_json=new_json,
            graph_old=None,
            graph_new=graph_new_json
        ))
    else:
        # Comparison upload
        old_text = extract_text_from_pdf(upload.old_path)
        new_text = extract_text_from_pdf(upload.new_path)

        old_summary = get_summary_with_context(old_text)
        new_summary = get_summary_with_context(new_text, context=old_summary)

        old_json = get_entity_relationship_with_context(old_summary)
        new_json = get_entity_relationship_with_context(new_summary, context=old_json)

        G_old = parse_graph_data(json.loads(old_json))
        G_new = parse_graph_data(json.loads(new_json))

        graph_old_json = json.dumps({
            "nodes": [{"id": n, **G_old.nodes[n]} for n in G_old.nodes],
            "edges": [{"from": u, "to": v, **G_old[u][v]} for u, v in G_old.edges]
        })

        graph_new_json = json.dumps({
            "nodes": [{"id": n, **G_new.nodes[n]} for n in G_new.nodes],
            "edges": [{"from": u, "to": v, **G_new[u][v]} for u, v in G_new.edges]
        })

        db.session.add(Summary(upload_id=upload.id, old_summary=old_summary, new_summary=new_summary))
        db.session.add(EntityGraph(
            upload_id=upload.id,
            old_json=old_json,
            new_json=new_json,
            graph_old=graph_old_json,
            graph_new=graph_new_json
        ))

    db.session.commit()


@app.route("/compare/<int:upload_id>")
def compare(upload_id):
    return render_template("compare.html", upload_id=upload_id)

@app.route("/graph_data/<int:upload_id>/<version>")
def graph_data(upload_id, version):
    graph_entry = EntityGraph.query.filter_by(upload_id=upload_id).first()
    if not graph_entry:
        return jsonify({"nodes": [], "edges": []})
    if version == "old":
        return jsonify(json.loads(graph_entry.graph_old or '{}'))
    if version == "new":
        return jsonify(json.loads(graph_entry.graph_new or '{}'))
    return jsonify({"error": "Invalid version"}), 400

@app.route("/regenerate/<int:upload_id>", methods=["POST"])
def regenerate(upload_id):
    process_upload(upload_id)
    return redirect(url_for("compare", upload_id=upload_id))

@app.route("/approve/<int:upload_id>", methods=["POST"])
def approve(upload_id):
    summary = Summary.query.filter_by(upload_id=upload_id).first()
    graph = EntityGraph.query.filter_by(upload_id=upload_id).first()
    if not summary or not graph:
        return "Data not found", 404

    if not summary.new_summary or not graph.new_json:
        return "New data missing. Please upload new regulation first.", 400

    prompt = (
        "Given the original document, pickup the modality of reporting. "
        "Given this particular graph, pickup the necessary actions to be performed. "
        "Generate a KOP document with step wise instruction for operational personel.\n\n"
        f"Summary:\n{summary.new_summary}\n\n"
        f"Entity Relationship JSON:\n{graph.new_json}"
    )

    kop_response = get_summary_with_context(prompt)

    doc = Document()
    doc.add_heading("Key Operating Procedure (KOP)", 0)
    markdown_to_docx(doc, kop_response)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"kop_upload_{upload_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/history")
def history():
    uploads = Upload.query.order_by(Upload.created_at.desc()).all()
    return render_template("history.html", uploads=uploads)

if __name__ == "__main__":
    with app.app_context():
        db.create_all()
        if not Regulation.query.first():
            db.session.add_all([
                Regulation(name="EMIR Refit"),
                Regulation(name="MiFID II"),
                Regulation(name="SFTR")
            ])
            db.session.commit()
    app.run(debug=True)
