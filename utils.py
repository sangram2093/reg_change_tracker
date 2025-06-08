import fitz  # PyMuPDF
import json
import networkx as nx
from docx import Document
import re

def extract_text_from_pdf(path):
    """Extracts text from a PDF file."""
    text = ""
    with fitz.open(path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def parse_graph_data(relationship_json):
    """Parses entity relationship JSON into NetworkX DiGraph with vis.js-compatible properties."""
    if isinstance(relationship_json, str):
        relationship_json = json.loads(relationship_json)

    G = nx.DiGraph()

    for entity in relationship_json.get("entities", []):
        G.add_node(entity["id"], label=entity["name"], group=entity["type"])

    for rel in relationship_json.get("relationships", []):
        tooltip_lines = [
            f"Verb: {rel.get('verb', '')}",
            f"Optionality: {rel.get('Optionality', '')}",
            f"Condition: {rel.get('Condition for Relationship to be Active', '')}",
            f"Property: {rel.get('Property of Object (part of condition)', '')}",
            f"Thresholds: {rel.get('Thresholds', '')}",
            f"Frequency: {rel.get('frequency', '')}"
        ]
        tooltip_text = "\n".join(tooltip_lines)

        G.add_edge(
            rel["subject_id"],
            rel["object_id"],
            label=rel.get("verb", ""),
            title=tooltip_text
        )

    return G

def compare_graphs(G_old, G_new):
    """Compares old and new graphs to identify changed edges and nodes."""
    changed_edges = []
    added_nodes = list(set(G_new.nodes) - set(G_old.nodes))
    removed_nodes = list(set(G_old.nodes) - set(G_new.nodes))

    for u, v in G_new.edges:
        if G_old.has_edge(u, v):
            old_data = G_old[u][v]
            new_data = G_new[u][v]
            if old_data.get("label") != new_data.get("label") or old_data.get("title") != new_data.get("title"):
                changed_edges.append((u, v))
        else:
            changed_edges.append((u, v))

    return changed_edges, added_nodes, removed_nodes

def markdown_to_docx(doc: Document, text: str):
    """Converts markdown-style Gemini response into formatted Word content."""
    lines = text.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # Headings
        if stripped.startswith("###"):
            doc.add_heading(stripped.lstrip("#").strip(), level=3)
        elif stripped.startswith("##"):
            doc.add_heading(stripped.lstrip("#").strip(), level=2)
        elif stripped.startswith("#"):
            doc.add_heading(stripped.lstrip("#").strip(), level=1)

        # Bullets
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style='List Bullet')

        # Numbered list
        elif re.match(r'^\d+\.\s', stripped):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', stripped), style='List Number')

        # Bold
        elif "**" in stripped:
            para = doc.add_paragraph()
            while "**" in stripped:
                if "**" not in stripped:
                    para.add_run(stripped)
                    break
                before, bold, rest = stripped.split("**", 2)
                para.add_run(before)
                bold_run = para.add_run(bold)
                bold_run.bold = True
                stripped = rest
            para.add_run(stripped)

        else:
            doc.add_paragraph(stripped)

def generate_kop_docx(new_summary, new_json_str, output_path, llm_client):
    """Generates a Word doc KOP using Gemini/Vertex LLM output."""
    prompt = (
        "Given the original document, pickup the modality of reporting. "
        "Given this particular graph, pickup the necessary actions to be performed. "
        "Generate a KOP document with step wise instruction for operational personnel."
    )

    entity_data = json.loads(new_json_str)
    full_prompt = f"{prompt}\n\nDocument Summary:\n{new_summary}\n\nEntity Relationships:\n{json.dumps(entity_data, indent=2)}"

    kop_response = llm_client(full_prompt)

    doc = Document()
    doc.add_heading("Key Operating Procedure (KOP)", level=1)
    markdown_to_docx(doc, kop_response)
    doc.save(output_path)
